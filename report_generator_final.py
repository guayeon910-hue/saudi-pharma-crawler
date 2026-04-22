"""
report_generator_final.py -- SA_최종 합본 DOCX 생성
표지 + SA_01 + SA_02 + SA_03을 섹션 브레이크로 병합한다.
"""

from __future__ import annotations

import copy
import logging
import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from _report_styles import (
    COLORS,
    FONTS,
    apply_page_margins,
    set_cell_shading,
)

logger = logging.getLogger("report_generator_final")

REPORTS_DIR = Path(__file__).resolve().parent / "reports"


def _add_page_break(doc: Document) -> None:
    """문서에 페이지 브레이크 단락을 추가한다."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _merge_docs(base: Document, *part_paths: Path) -> None:
    """각 파트 docx의 body 요소를 base 문서에 추가한다."""
    for path in part_paths:
        if not path or not path.exists():
            logger.warning("Part file missing: %s", path)
            continue
        _add_page_break(base)
        part = Document(str(path))
        for elem in part.element.body:
            # sectPr(섹션 속성)은 마지막 섹션만 건드리므로 건너뜀
            if elem.tag.endswith("}sectPr"):
                continue
            base.element.body.append(copy.deepcopy(elem))


def generate_final_report(
    p1_path: Path | None,
    p2_path: Path | None,
    p3_path: Path | None,
    meta: dict | None = None,
    output_dir: Path | None = None,
) -> Path:
    """SA_최종 합본 DOCX 생성."""
    out_dir = output_dir or REPORTS_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    meta      = meta or {}
    company   = os.environ.get("REPORT_COMPANY_NAME", "KOREA UNITED PHARM INC.")
    now       = datetime.now(timezone.utc)
    date_str  = now.strftime("%Y%m%d")
    date_line = now.strftime("%Y-%m-%d")
    timestamp = now.strftime("%Y-%m-%d %H:%M UTC")

    trade_name  = meta.get("trade_name", "")
    inn         = meta.get("inn") or meta.get("ingredient", "")
    hs_code     = meta.get("hs_code", "")
    dosage_form = meta.get("dosage_form", "")
    strength    = meta.get("strength", "")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONTS["ko"]
    style.font.size = Pt(10)

    # 표지는 3.17 cm 여백
    section = doc.sections[0]
    apply_page_margins(section, cover=True)

    # ── 표지 (SG_00 스타일) ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(company)
    cr.bold = True
    cr.font.size = Pt(11)
    cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr_ = tp.add_run("사우디아라비아 수출 종합 보고서")
    tr_.bold = True
    tr_.font.size = Pt(28)
    tr_.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    if trade_name or inn:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_text = trade_name + (f" — {inn}" if inn else "")
        sr_ = sp.add_run(sub_text)
        sr_.font.size = Pt(18)
        sr_.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(date_line)
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    # 제품 상세 박스
    if any([hs_code, dosage_form, strength]):
        detail_tbl = doc.add_table(rows=1, cols=1)
        detail_tbl.style = "Table Grid"
        dcell = detail_tbl.rows[0].cells[0]
        set_cell_shading(dcell, COLORS["light_gray"])
        dcell.text = ""
        detail_lines = []
        if hs_code:
            detail_lines.append(f"HS 코드: {hs_code}")
        if dosage_form:
            detail_lines.append(f"제형: {dosage_form}")
        if strength:
            detail_lines.append(f"함량: {strength}")
        dp2 = dcell.paragraphs[0]
        dp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr2 = dp2.add_run("  |  ".join(detail_lines))
        dr2.font.size = Pt(12)
        dr2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph("")

    # 목차
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_r = toc_p.add_run("목 차")
    toc_r.bold = True
    toc_r.font.size = Pt(14)
    toc_r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    toc_items = [
        ("SA_01", "사우디아라비아 시장 분석 보고서"),
        ("SA_02", "수출가격전략 보고서"),
        ("SA_03", "바이어 발굴 보고서"),
    ]
    for code, title in toc_items:
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ir = ip.add_run(f"{code}  {title}")
        ir.font.size = Pt(12)
        ir.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # 각 파트 병합
    _merge_docs(doc, p1_path, p2_path, p3_path)

    # 최종 푸터
    _add_page_break(doc)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"Saudi Pharma Crawler | {timestamp}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    filename = f"sa_final_{date_str}.docx"
    output_path = out_dir / filename
    doc.save(str(output_path))
    logger.info("Final report saved: %s", output_path)
    return output_path
