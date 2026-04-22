"""
_report_styles.py -- SG 양식 기준 공통 스타일 헬퍼
모든 SA 보고서 generator가 import해서 사용한다.
"""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

COLORS = {
    "navy":      "1B3A6B",
    "orange":    "C85A00",
    "blue":      "1457A0",
    "green":     "1A6B35",
    "red":       "C0392B",
    "gray_text": "888888",
    "body":      "1A1A1A",
    "white":     "FFFFFF",
    "light_gray":"EEEEEE",
}

FONTS = {
    "ko":  "Malgun Gothic",
    "num": "Helvetica Neue",
}

_CM_BODY_TOP   = 2.26
_CM_BODY_BOT   = 2.54
_CM_BODY_SIDE  = 2.26
_CM_COVER      = 3.17


def set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def apply_page_margins(section, cover: bool = False) -> None:
    if cover:
        m = Cm(_CM_COVER)
        section.top_margin = m
        section.bottom_margin = m
        section.left_margin = m
        section.right_margin = m
    else:
        section.top_margin    = Cm(_CM_BODY_TOP)
        section.bottom_margin = Cm(_CM_BODY_BOT)
        section.left_margin   = Cm(_CM_BODY_SIDE)
        section.right_margin  = Cm(_CM_BODY_SIDE)


def set_header_footer(section, company: str, product: str, date: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    # 기존 단락 초기화
    for p in header.paragraphs:
        p.clear()
    if not header.paragraphs:
        hp = header.add_paragraph()
    else:
        hp = header.paragraphs[0]
    hp.clear()
    left_run = hp.add_run(company)
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    # 탭으로 오른쪽 정렬
    hp.add_run("\t")
    right_run = hp.add_run(f"{product}  {date}")
    right_run.font.size = Pt(9)
    right_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        fp = footer.add_paragraph()
    else:
        fp = footer.paragraphs[0]
    fp.clear()
    # 페이지 번호
    fp.add_run("\t")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_elem = OxmlElement("w:r")
    run_elem.append(fld_begin)
    run_elem.append(instrText)
    run_elem.append(fld_sep)
    run_elem.append(fld_end)
    fp._p.append(run_elem)


def section_header_row(table, text: str, col_span: int) -> None:
    """테이블 첫 행: 1B3A6B 배경 + 흰 글씨 + 14pt Bold 섹션 제목."""
    row = table.rows[0]
    merged = row.cells[0]
    if col_span > 1:
        merged = row.cells[0].merge(row.cells[col_span - 1])
    merged.text = ""
    set_cell_shading(merged, COLORS["navy"])
    p = merged.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = FONTS["ko"]


def colored_price_box(doc, label: str, usd: str, sar: str, krw: str, color_hex: str) -> None:
    """3-컬럼 가격 박스 (저가/기준/프리미엄)."""
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    # 헤더 셀 병합
    hdr = tbl.rows[0].cells[0].merge(tbl.rows[0].cells[2])
    hdr.text = ""
    set_cell_shading(hdr, color_hex)
    hp = hdr.paragraphs[0]
    hr = hp.add_run(label)
    hr.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hr.font.name = FONTS["ko"]
    # 값 행
    vals = [("USD", usd), ("SAR", sar), ("KRW", krw)]
    for i, (cur, val) in enumerate(vals):
        cell = tbl.rows[1].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{cur}\n")
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r2 = p.add_run(str(val))
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.name = FONTS["num"]


def verdict_cell(cell, verdict_ko: str) -> None:
    """판정 셀 색상: 적합=green, 조건부=orange, 부적합=red."""
    color_map = {
        "적합":   COLORS["green"],
        "조건부": COLORS["orange"],
        "부적합": COLORS["red"],
    }
    bg = color_map.get(verdict_ko, COLORS["light_gray"])
    set_cell_shading(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(verdict_ko)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if bg != COLORS["light_gray"] else RGBColor(0x1A, 0x1A, 0x1A)
    run.font.size = Pt(11)
