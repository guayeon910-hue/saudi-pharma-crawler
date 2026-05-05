"""
report_generator.py -- 한국어 시장 분석 보고서 DOCX 자동 생성

AggregatedResult + Claude 분석을 받아 스크린샷 템플릿에 맞춘 DOCX를 생성한다.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from _report_styles import (
    COLORS,
    FONTS,
    apply_page_margins,
    set_cell_shading as _set_cell_shading,
    set_header_footer,
    verdict_cell,
)
from report_refine import (
    aggregate_evidence_by_source,
    count_procurement_vs_retail,
    fill_pillar_fallbacks,
    fill_strategy_fallbacks,
    refine_cell_text,
)

logger = logging.getLogger("report_generator")

REPORTS_DIR = Path(__file__).resolve().parent / "reports"

# 사우디 파이프라인 참조 사이트 (카테고리)
REFERENCE_SITE_CATEGORIES_KSA: dict[str, list[str]] = {
    "공공조달": [
        "SFDA — Registered products / pricing references",
        "NUPCO / Government procurement (when applicable)",
        "Hospital formulary listings (KSA)",
    ],
    "규제·등재": [
        "SFDA Drug Register",
        "Saudi National Formulary (reference)",
        "WHO EML (cross-check)",
    ],
    "시장 거시": [
        "World Bank — Saudi Arabia health / macro",
        "KOTRA 무역관 — KSA 의약품",
        "SFDA / MOH public communications",
    ],
    "규제 프레임워크": [
        "Saudi FDA regulations & registration pathways",
        "GCC harmonization (where applicable)",
        "PubMed / clinical evidence (supporting dossiers)",
    ],
}



def _add_table_row(table, cells_data: list[str], bold_first: bool = False) -> None:
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        if bold_first and i == 0:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _verdict_display(verdict: str) -> str:
    v = (verdict or "").strip()
    legacy = {
        "가능": "적합",
        "possible": "적합",
        "conditional": "조건부",
        "불가": "부적합",
        "impossible": "부적합",
        "분석실패": "—",
        "API 키 미설정": "—",
    }
    if v in legacy:
        return legacy[v]
    if v in ("적합", "조건부", "부적합"):
        return v
    return v or "—"


def _domain_from_url(url: str) -> str:
    if not url:
        return ""
    try:
        u = urlparse(url)
        host = u.netloc or ""
        return host.replace("www.", "") or ""
    except Exception:
        return ""


def generate_report(
    drug: Any,
    search_data: dict,
    *,
    analysis: dict | None = None,
    refs: list[dict] | None = None,
    report_meta: dict | None = None,
    exchange_rates: dict | None = None,
) -> Path:
    """한국어 시장 분석 보고서 DOCX 생성."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    # 환율 변환 헬퍼 (SAR 기준)
    _sar_usd = (exchange_rates or {}).get("sar_usd") or (1 / 3.75)
    _sar_krw = (exchange_rates or {}).get("sar_krw") or 392.64
    _fx_source = (exchange_rates or {}).get("source", "fallback")

    def _to_usd(price_sar: Any) -> str:
        try:
            return f"${float(price_sar) * _sar_usd:.2f}"
        except (TypeError, ValueError):
            return ""

    def _to_krw(price_sar: Any) -> str:
        try:
            return f"₩{int(float(price_sar) * _sar_krw):,}"
        except (TypeError, ValueError):
            return ""

    def _price_sar(match: dict[str, Any]) -> Any:
        for key in ("price_sar", "price_local", "price", "retail_price"):
            value = match.get(key)
            if value not in (None, ""):
                return value
        return ""

    if isinstance(drug, dict):
        trade_name = drug.get("trade_name", "Unknown")
        ingredient = drug.get("ingredient", "")
        strength = drug.get("strength", "")
        dosage_form = drug.get("dosage_form", "")
        drug_type = drug.get("drug_type", "")
        drug_id = drug.get("id", "unknown")
    else:
        trade_name = drug.trade_name
        ingredient = drug.ingredient
        strength = drug.strength
        dosage_form = drug.dosage_form
        drug_type = drug.drug_type
        drug_id = drug.id

    refs = refs or []
    report_meta = report_meta or {}

    if analysis is None:
        analysis = {
            "verdict": search_data.get("export_feasibility", "분석실패"),
            "confidence": 0.0,
            "rationale": search_data.get("feasibility_rationale", ""),
            "key_factors": [],
            "hs_code": None,
            "case_type": None,
            "pillars": {},
            "strategy": {},
        }

    pillars = fill_pillar_fallbacks(analysis)
    strategy = fill_strategy_fallbacks(analysis)

    for k in list(pillars.keys()):
        pillars[k] = refine_cell_text(pillars[k])
    for k in list(strategy.keys()):
        strategy[k] = refine_cell_text(strategy[k])

    verdict_ko = _verdict_display(analysis.get("verdict", ""))
    conf = float(analysis.get("confidence") or 0.0)
    hs = analysis.get("hs_code")
    hs_s = hs.strip() if isinstance(hs, str) else ""
    case_t = analysis.get("case_type")
    case_s = case_t.strip() if isinstance(case_t, str) else ""

    company = os.environ.get("REPORT_COMPANY_NAME", "KOREA UNITED PHARM INC.")
    report_title = os.environ.get("REPORT_TITLE", "사우디아라비아 시장 분석 보고서")

    now = datetime.now(timezone.utc)
    date_line = report_meta.get("collection_finished_at") or now.strftime("%Y-%m-%d")
    timestamp = now.strftime("%Y-%m-%d %H:%M UTC")
    date_str = now.strftime("%Y%m%d")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONTS["ko"]
    style.font.size = Pt(10)

    # 페이지 여백 (SG 본문 양식)
    section = doc.sections[0]
    apply_page_margins(section, cover=False)
    set_header_footer(section, company, trade_name, date_line)

    # ── 표지 헤더 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(company)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(report_title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub_p.add_run(f"{trade_name} — {ingredient}")
    sr2.font.size = Pt(18)
    sr2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(date_line)
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 제품 상세 정보줄 (메타 바 대체)
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_parts = []
    if strength or dosage_form:
        meta_parts.append(f"{strength} {dosage_form}".strip())
    if hs_s:
        meta_parts.append(f"HS {hs_s}")
    if case_s:
        meta_parts.append(case_s)
    if conf:
        meta_parts.append(f"신뢰도 {conf:.0%}")
    mr2 = meta_p.add_run("  |  ".join(meta_parts))
    mr2.font.size = Pt(12)
    mr2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(6)
    p_sep = sep.add_run("—" * 42)
    p_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph("")

    # 1. 진출 적합 판정
    sec1 = doc.add_table(rows=2, cols=2)
    sec1.style = "Table Grid"
    h1 = sec1.rows[0].cells[0].merge(sec1.rows[0].cells[1])
    h1.text = ""
    hp1 = h1.paragraphs[0]
    r_h1 = hp1.add_run("1. 진출 적합 판정")
    r_h1.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_shading(h1, COLORS["navy"])
    sec1.rows[1].cells[0].text = "판정"
    verdict_cell(sec1.rows[1].cells[1], verdict_ko)
    for para in sec1.rows[1].cells[0].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
    doc.add_paragraph("")

    # 2. 판정 근거 (1~5)
    labels = (
        ("1", "시장·의료", "market_medical"),
        ("2", "규제", "regulation"),
        ("3", "무역", "trade"),
        ("4", "조달", "procurement"),
        ("5", "유통", "distribution"),
    )
    t2 = doc.add_table(rows=1 + len(labels), cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].merge(t2.rows[0].cells[2])
    t2.rows[0].cells[0].text = ""
    r_t2 = t2.rows[0].cells[0].paragraphs[0].add_run("2. 판정 근거")
    r_t2.bold = True
    r_t2.font.size = Pt(14)
    r_t2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_shading(t2.rows[0].cells[0], COLORS["navy"])

    for i, (num, label, key) in enumerate(labels, start=1):
        row = t2.rows[i]
        row.cells[0].text = num
        row.cells[1].text = label
        row.cells[2].text = pillars.get(key, "")
    doc.add_paragraph("")

    # 3. 시장 진출 전략
    strat_rows = [
        ("진입 채널 전략", strategy.get("entry_channels", "")),
        ("가격 포지셔닝", strategy.get("price_positioning", "")),
        ("유통 파트너", strategy.get("distribution_partners", "")),
        ("리스크·조건", strategy.get("risk_conditions", "")),
    ]
    t3 = doc.add_table(rows=1 + len(strat_rows), cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].merge(t3.rows[0].cells[1])
    t3.rows[0].cells[0].text = ""
    sp = t3.rows[0].cells[0].paragraphs[0]
    r_t3 = sp.add_run("3. 시장 진출 전략")
    r_t3.bold = True
    r_t3.font.size = Pt(14)
    r_t3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_shading(t3.rows[0].cells[0], COLORS["navy"])

    for i, (title, body) in enumerate(strat_rows, start=1):
        t3.rows[i].cells[0].text = title
        t3.rows[i].cells[1].text = body

    doc.add_paragraph("")
    source_results = search_data.get("source_results") or []
    pub_n, priv_n = count_procurement_vs_retail(
        source_results if isinstance(source_results, list) else []
    )
    note_p = doc.add_paragraph()
    note_p.add_run(
        f"(참고) 공공·민간 표본: 낙찰 {pub_n}건 · 민간 {priv_n}건"
    ).italic = True
    doc.add_paragraph("")

    # 4. 근거 및 출처
    doc.add_paragraph()
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. 근거 및 출처")
    r_h4.bold = True
    r_h4.font.size = Pt(14)
    r_h4.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    sub = doc.add_paragraph()
    r_sub = sub.add_run("참조 데이터")
    r_sub.bold = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    ev_rows = aggregate_evidence_by_source(
        source_results if isinstance(source_results, list) else []
    )
    ev_table = doc.add_table(rows=1 + max(1, len(ev_rows)), cols=3)
    ev_table.style = "Table Grid"
    ev_headers = ["출처", "건수", "신뢰도(평균)"]
    for i, h in enumerate(ev_headers):
        c = ev_table.rows[0].cells[i]
        c.text = ""
        r_ev = c.paragraphs[0].add_run(h)
        r_ev.bold = True
        r_ev.font.size = Pt(10)
        r_ev.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(c, COLORS["navy"])
    if ev_rows:
        for j, er in enumerate(ev_rows, start=1):
            ev_table.rows[j].cells[0].text = str(er["source"])
            ev_table.rows[j].cells[1].text = f"{er['count']}건"
            ev_table.rows[j].cells[2].text = f"{er['avg_confidence']:.2f}"
    else:
        ev_table.rows[1].cells[0].text = "—"
        ev_table.rows[1].cells[1].text = "0건"
        ev_table.rows[1].cells[2].text = "—"

    doc.add_paragraph("")

    # 참조 사이트 (카테고리)
    rs_p = doc.add_paragraph()
    rs_p.add_run("참조 사이트 (카테고리)").bold = True
    for cat, sites in REFERENCE_SITE_CATEGORIES_KSA.items():
        doc.add_paragraph(f"▸ {cat}: " + ", ".join(sites))

    doc.add_paragraph("")

    # 수집 메타데이터 박스
    meta_box = doc.add_table(rows=1, cols=1)
    meta_box.style = "Table Grid"
    mcell = meta_box.rows[0].cells[0]
    _set_cell_shading(mcell, "EEEEEE")
    mcell.text = ""
    mp2 = mcell.paragraphs[0]
    lines = [
        f"최종 수집: {date_line}",
        f"수집 방식: {report_meta.get('collection_method', '—')}",
        f"의미적 신선도 판정: {report_meta.get('freshness_note', '—')}",
        f"LLM 본문 생성: {report_meta.get('llm_body_note', '—')}",
    ]
    sd = report_meta.get("search_duration_sec")
    if sd:
        lines.append(f"검색 소요: {sd:.1f}초")
    mp2.add_run("\n".join(lines)).font.size = Pt(9)

    doc.add_paragraph("")

    # Perplexity 추천 논문
    pp_p = doc.add_paragraph()
    pp_p.add_run("Perplexity 추천 논문").bold = True
    cache_note = ""
    for r0 in refs:
        if r0.get("cache_hit"):
            cache_note = " (cache_hit)"
            break
    doc.add_paragraph(f"상태{cache_note}".strip())

    if refs:
        for idx, ref in enumerate(refs[:15], start=1):
            title = str(ref.get("title") or ref.get("url") or "Untitled")
            url = str(ref.get("url") or "")
            dom = _domain_from_url(url)
            date_part = ref.get("published") or ref.get("date") or ""
            tail = f" ({dom}" + (f", {date_part}" if date_part else "") + ")"
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Cm(0.5)
            if url:
                line.add_run(f"{idx}. {title}{tail}\n   {url}")
            else:
                line.add_run(f"{idx}. {title}{tail}")
    else:
        doc.add_paragraph("Perplexity 참고 링크가 없습니다(API 키 또는 검색 결과 없음).")

    doc.add_paragraph("")

    # 부록: 요약 경쟁 데이터 (짧게)
    doc.add_paragraph().add_run("부록: 수집 요약").bold = True
    total_matches = search_data.get("total_matches", 0)
    doc.add_paragraph(f"총 매칭 건수: {total_matches} | 생성 시각(UTC): {timestamp}")
    fx_note = doc.add_paragraph(
        f"※ 가격 추정치 안내: SAR 기준 공개 데이터 역산값으로 실제 입찰·협상가와 차이가 있을 수 있습니다. "
        f"USD/KRW 환산은 {_fx_source} 환율(1 SAR ≈ ${_sar_usd:.4f} / ₩{_sar_krw:,.0f}) 기준입니다. "
        "의사결정 전 현지 에이전트 및 최신 환율을 확인하시기 바랍니다."
    )
    fx_note.runs[0].font.size = Pt(8)
    fx_note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    sfda_matches = []
    for sr in source_results if isinstance(source_results, list) else []:
        if sr.get("source_category") == "공공조달" and sr.get("matches"):
            sfda_matches.extend(sr["matches"])
    if sfda_matches:
        doc.add_paragraph("SFDA/공공조달 샘플 (최대 8건):")
        ap = doc.add_table(rows=1 + min(8, len(sfda_matches)), cols=6)
        ap.style = "Table Grid"
        for i, h in enumerate(["품목", "성분", "함량", "SAR", "USD", "KRW"]):
            ap.rows[0].cells[i].text = h
        for ri, m in enumerate(sfda_matches[:8], start=1):
            p_sar = _price_sar(m)
            ap.rows[ri].cells[0].text = str(m.get("trade_name", ""))
            ap.rows[ri].cells[1].text = str(m.get("scientific_name", ""))
            ap.rows[ri].cells[2].text = str(m.get("strength", ""))
            ap.rows[ri].cells[3].text = str(p_sar)
            ap.rows[ri].cells[4].text = _to_usd(p_sar)
            ap.rows[ri].cells[5].text = _to_krw(p_sar)

    retail_matches = []
    for sr in source_results if isinstance(source_results, list) else []:
        if sr.get("source_category") == "민간" and sr.get("matches"):
            retail_matches.extend(
                [{**m, "_source": sr.get("source_name", "")} for m in sr["matches"][:8]]
            )
    if retail_matches:
        doc.add_paragraph("민간 소매 샘플 (최대 8건):")
        rp = doc.add_table(rows=1 + min(8, len(retail_matches)), cols=5)
        rp.style = "Table Grid"
        for i, h in enumerate(["상품", "SAR", "USD", "KRW", "출처"]):
            rp.rows[0].cells[i].text = h
        for ri, m in enumerate(retail_matches[:8], start=1):
            p_sar = _price_sar(m)
            rp.rows[ri].cells[0].text = str(m.get("name", m.get("trade_name", "")))
            rp.rows[ri].cells[1].text = str(p_sar)
            rp.rows[ri].cells[2].text = _to_usd(p_sar)
            rp.rows[ri].cells[3].text = _to_krw(p_sar)
            rp.rows[ri].cells[4].text = str(m.get("_source", ""))[:40]

    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"Saudi Pharma Crawler | {timestamp}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    filename = f"market_report_{drug_id}_{date_str}.docx"
    output_path = REPORTS_DIR / filename
    doc.save(str(output_path))
    logger.info("Report saved: %s", output_path)
    return output_path


if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(message)s")

    sys.path.insert(0, str(Path(__file__).resolve().parent / "assets" / "snippets"))
    from drug_registry import DrugRegistry
    from targeted_search import search_one_drug

    reg = DrugRegistry()
    drugs = reg.list_drugs()

    if not drugs:
        print("No drugs registered.")
        sys.exit(1)

    drug_id = sys.argv[1] if len(sys.argv) > 1 else drugs[0].id
    drug = reg.get_drug(drug_id)
    if not drug:
        print(f"Drug '{drug_id}' not found.")
        sys.exit(1)

    print(f"Searching: {drug.trade_name}...")
    result = search_one_drug(drug)
    search_data = result.to_dict()

    print("Generating report...")
    path = generate_report(drug, search_data)
    print(f"Report saved: {path}")
