"""
report_generator_p2.py -- SA_02 수출가격전략 보고서 DOCX 생성
/api/p2/price-analyze 응답 데이터를 받아 SG 양식으로 보고서를 생성한다.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from _report_styles import (
    COLORS,
    FONTS,
    apply_page_margins,
    colored_price_box,
    set_cell_shading,
    set_header_footer,
)

logger = logging.getLogger("report_generator_p2")

REPORTS_DIR = Path(__file__).resolve().parent / "reports"


def _sec_hdr(doc, text: str) -> None:
    """네이비 배경 흰 글씨 섹션 헤더 단락."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, COLORS["navy"])
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = FONTS["ko"]
    doc.add_paragraph("")


def _sub_hdr(doc, text: str) -> None:
    """12pt Bold 네이비 소섹션 헤더."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    run.font.name = FONTS["ko"]


def generate_p2_report(data: dict, output_dir: Path | None = None) -> Path:
    """SA_02 수출가격전략 DOCX 보고서 생성."""
    out_dir = output_dir or REPORTS_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    company     = os.environ.get("REPORT_COMPANY_NAME", "KOREA UNITED PHARM INC.")
    now         = datetime.now(timezone.utc)
    date_str    = now.strftime("%Y%m%d")
    date_line   = now.strftime("%Y-%m-%d")
    timestamp   = now.strftime("%Y-%m-%d %H:%M UTC")

    product     = data.get("product", {}) or {}
    trade_name  = product.get("trade_name") or data.get("trade_name", "Unknown")
    inn         = product.get("inn") or product.get("ingredient", "")

    scenarios_raw    = data.get("scenarios", []) or []
    if isinstance(scenarios_raw, dict):
        scenarios = []
        for key, value in scenarios_raw.items():
            if isinstance(value, dict):
                scenarios.append({"name": key, **value})
    else:
        scenarios = [s for s in scenarios_raw if isinstance(s, dict)]
    competitor_stats = data.get("competitor_stats", {}) or {}
    regulatory_cost  = data.get("regulatory_cost", {}) or {}
    exchange_rates   = data.get("exchange_rates", {}) or {}
    classification   = data.get("classification", {}) or {}
    notes            = data.get("notes", []) or []
    price_pool_sources = data.get("price_pool_sources", []) or []

    sar_usd = exchange_rates.get("sar_usd") or (1 / 3.75)
    sar_krw = exchange_rates.get("sar_krw") or 392.64
    fx_src  = exchange_rates.get("source", "fallback")

    def _fmt_usd(v):
        try:
            return f"${float(v):.2f}"
        except Exception:
            return str(v) if v else "—"

    def _fmt_sar(v):
        try:
            return f"SAR {float(v):.2f}"
        except Exception:
            return str(v) if v else "—"

    def _fmt_krw(v):
        try:
            return f"₩{int(float(v)):,}"
        except Exception:
            return str(v) if v else "—"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONTS["ko"]
    style.font.size = Pt(10)

    section = doc.sections[0]
    apply_page_margins(section, cover=False)
    set_header_footer(section, company, trade_name, date_line)

    # ── 표지줄 ──
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(company)
    cr.bold = True
    cr.font.size = Pt(11)
    cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr_ = tp.add_run("수출가격전략 보고서")
    tr_.bold = True
    tr_.font.size = Pt(28)
    tr_.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr_ = sp.add_run(f"{trade_name}" + (f" — {inn}" if inn else ""))
    sr_.font.size = Pt(18)
    sr_.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(date_line)
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    # ── 섹션 1: 시장 가격 맥락 ──
    _sec_hdr(doc, "1. 시장 가격 맥락")
    _sub_hdr(doc, "경쟁 제품 가격 분포 (SAR)")

    stat_keys = [
        ("min",    "최솟값"),
        ("p25",    "25th 백분위"),
        ("median", "중앙값"),
        ("p75",    "75th 백분위"),
        ("max",    "최댓값"),
        ("count",  "샘플 수"),
    ]
    has_stats = any(competitor_stats.get(k) is not None for k, _ in stat_keys)
    if has_stats:
        stat_tbl = doc.add_table(rows=1 + len(stat_keys), cols=3)
        stat_tbl.style = "Table Grid"
        for i, h in enumerate(["지표", "SAR", "USD"]):
            c = stat_tbl.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(c, COLORS["navy"])
        for i, (key, label) in enumerate(stat_keys, start=1):
            val = competitor_stats.get(key)
            stat_tbl.rows[i].cells[0].text = label
            stat_tbl.rows[i].cells[1].text = _fmt_sar(val) if key != "count" else (str(val) if val else "—")
            stat_tbl.rows[i].cells[2].text = _fmt_usd(float(val) * sar_usd) if key != "count" and val else "—"
    else:
        doc.add_paragraph("경쟁 제품 가격 데이터 없음.")

    doc.add_paragraph("")

    _sub_hdr(doc, "가격 풀 구성")
    if price_pool_sources:
        ps_tbl = doc.add_table(rows=1 + len(price_pool_sources), cols=3)
        ps_tbl.style = "Table Grid"
        for i, h in enumerate(["출처", "유형", "샘플 수"]):
            c = ps_tbl.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(c, COLORS["navy"])
        for j, src in enumerate(price_pool_sources, start=1):
            ps_tbl.rows[j].cells[0].text = str(src.get("name") or src.get("source") or "—")
            ps_tbl.rows[j].cells[1].text = str(src.get("type") or src.get("origin") or "—")
            ps_tbl.rows[j].cells[2].text = str(src.get("count") or src.get("sample_count") or "—")
    else:
        doc.add_paragraph("가격 풀 출처 정보 없음.")

    doc.add_paragraph("")

    # ── 섹션 2: FOB 가격 시나리오 ──
    _sec_hdr(doc, "2. FOB 가격 시나리오")

    scenario_map = {
        "저가":   COLORS["orange"],
        "기준":   COLORS["blue"],
        "premium": COLORS["green"],
        "프리미엄": COLORS["green"],
    }

    def _get_fob(sc: dict, key: str):
        return sc.get(f"fob_{key}") or sc.get(key) or sc.get("fob") or ""

    for sc in scenarios:
        label = sc.get("label") or sc.get("scenario") or sc.get("name") or "시나리오"
        # label 정규화
        label_ko = label
        color = COLORS["blue"]
        for k, c in scenario_map.items():
            if k in label.lower() or k in label:
                color = c
                label_ko = {"저가": "저가 진입", "기준": "기준가", "premium": "프리미엄", "프리미엄": "프리미엄"}.get(k, label)
                break

        usd_v = _get_fob(sc, "usd")
        sar_v = _get_fob(sc, "sar")
        krw_v = _get_fob(sc, "krw")
        colored_price_box(doc, label_ko, _fmt_usd(usd_v), _fmt_sar(sar_v), _fmt_krw(krw_v), color)
        doc.add_paragraph("")

    if not scenarios:
        doc.add_paragraph("시나리오 데이터 없음.")
        doc.add_paragraph("")

    # ── 섹션 3: 역산 단계별 계산 ──
    _sec_hdr(doc, "3. 역산 단계별 계산")

    for sc in scenarios:
        label = sc.get("label") or sc.get("scenario") or sc.get("name") or "시나리오"
        _sub_hdr(doc, label)
        steps = sc.get("steps") or sc.get("calculation_steps") or []
        if steps:
            step_tbl = doc.add_table(rows=1 + len(steps), cols=3)
            step_tbl.style = "Table Grid"
            for i, h in enumerate(["단계", "항목", "값 (SAR)"]):
                c = step_tbl.rows[0].cells[i]
                c.text = ""
                r = c.paragraphs[0].add_run(h)
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(c, COLORS["navy"])
            for si, step in enumerate(steps, start=1):
                step_tbl.rows[si].cells[0].text = str(si)
                step_tbl.rows[si].cells[1].text = str(step.get("label") or step.get("name") or step.get("step") or "—")
                val = step.get("value_sar") or step.get("value") or step.get("amount") or ""
                step_tbl.rows[si].cells[2].text = _fmt_sar(val) if val else "—"
        else:
            doc.add_paragraph("역산 데이터 없음.")
        doc.add_paragraph("")

    if not scenarios:
        doc.add_paragraph("")

    # ── 섹션 4: 규제비 감가상각 ──
    _sec_hdr(doc, "4. 규제비 감가상각")

    reg_rows = [
        ("SFDA 등록비",      regulatory_cost.get("sfda_registration_sar") or regulatory_cost.get("sfda_registration") or regulatory_cost.get("sfda_fee")),
        ("SABER PCoC",       regulatory_cost.get("saber_pcoc_annual_sar") or regulatory_cost.get("saber_pcoc") or regulatory_cost.get("pcoc")),
        ("SCoC",             regulatory_cost.get("saber_scoc_annual_sar") or regulatory_cost.get("scoc")),
        ("연 수량 (units)",   (regulatory_cost.get("assumptions") or {}).get("annual_units") or regulatory_cost.get("annual_units") or regulatory_cost.get("annual_quantity")),
        ("유닛당 상각액 (SAR)", regulatory_cost.get("per_unit_amortization_sar") or regulatory_cost.get("per_unit_usd") or regulatory_cost.get("amortization_per_unit")),
    ]

    reg_tbl = doc.add_table(rows=1 + len(reg_rows), cols=2)
    reg_tbl.style = "Table Grid"
    for i, h in enumerate(["항목", "값"]):
        c = reg_tbl.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, COLORS["navy"])
    for i, (label, val) in enumerate(reg_rows, start=1):
        reg_tbl.rows[i].cells[0].text = label
        reg_tbl.rows[i].cells[1].text = str(val) if val is not None else "—"

    doc.add_paragraph("")

    # ── 섹션 5: 분류 및 전략 근거 ──
    _sec_hdr(doc, "5. 분류 및 전략 근거")

    cls_rows = [
        ("제품 유형",          classification.get("product_kind") or classification.get("type")),
        ("복합 제제 여부",      "복합" if classification.get("is_combination") else "단일"),
        ("전략 근거",          classification.get("rationale")),
    ]
    warnings = classification.get("warnings") or []

    cls_tbl = doc.add_table(rows=1 + len(cls_rows), cols=2)
    cls_tbl.style = "Table Grid"
    for i, h in enumerate(["항목", "내용"]):
        c = cls_tbl.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c, COLORS["navy"])
    for i, (label, val) in enumerate(cls_rows, start=1):
        cls_tbl.rows[i].cells[0].text = label
        cls_tbl.rows[i].cells[1].text = str(val) if val else "—"

    if warnings:
        doc.add_paragraph("")
        _sub_hdr(doc, "주의사항")
        if isinstance(warnings, list):
            for w in warnings:
                doc.add_paragraph(f"• {w}")
        else:
            doc.add_paragraph(str(warnings))

    doc.add_paragraph("")

    # ── 섹션 6: 출처 및 면책 ──
    _sec_hdr(doc, "6. 출처 및 면책")

    if price_pool_sources:
        _sub_hdr(doc, "가격 풀 출처")
        src_tbl = doc.add_table(rows=1 + len(price_pool_sources), cols=3)
        src_tbl.style = "Table Grid"
        for i, h in enumerate(["출처명", "유형", "비고"]):
            c = src_tbl.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(c, COLORS["navy"])
        for j, src in enumerate(price_pool_sources, start=1):
            src_tbl.rows[j].cells[0].text = str(src.get("name") or src.get("source") or "—")
            src_tbl.rows[j].cells[1].text = str(src.get("type") or src.get("origin") or "—")
            src_tbl.rows[j].cells[2].text = str(src.get("note") or src.get("url") or "—")
        doc.add_paragraph("")

    fx_p = doc.add_paragraph(
        f"※ 가격 추정치 안내: 공개 데이터 역산값으로 실제 입찰·협상가와 차이가 있을 수 있습니다. "
        f"USD/KRW 환산은 {fx_src} 환율(1 SAR ≈ ${sar_usd:.4f} / ₩{sar_krw:,.0f}) 기준입니다. "
        "의사결정 전 현지 에이전트 및 최신 환율을 확인하시기 바랍니다."
    )
    fx_p.runs[0].font.size = Pt(8)
    fx_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"Saudi Pharma Crawler | {timestamp}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    filename = f"sa_02_price_{date_str}.docx"
    output_path = out_dir / filename
    doc.save(str(output_path))
    logger.info("P2 report saved: %s", output_path)
    return output_path
