"""
report_generator_p3.py -- SA_03 바이어 발굴 보고서 DOCX 생성
/api/buyers/result 응답 데이터를 받아 SG 양식으로 보고서를 생성한다.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from _report_styles import (
    COLORS,
    FONTS,
    apply_page_margins,
    set_cell_shading,
    set_header_footer,
)

logger = logging.getLogger("report_generator_p3")

REPORTS_DIR = Path(__file__).resolve().parent / "reports"


def _sec_hdr(doc, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, COLORS["navy"])
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = FONTS["ko"]
    doc.add_paragraph("")


def _sub_hdr(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    run.font.name = FONTS["ko"]


def generate_p3_report(
    buyers: list,
    product_name: str,
    output_dir: Path | None = None,
) -> Path:
    """SA_03 바이어 발굴 DOCX 보고서 생성."""
    out_dir = output_dir or REPORTS_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    company   = os.environ.get("REPORT_COMPANY_NAME", "KOREA UNITED PHARM INC.")
    now       = datetime.now(timezone.utc)
    date_str  = now.strftime("%Y%m%d")
    date_line = now.strftime("%Y-%m-%d")
    timestamp = now.strftime("%Y-%m-%d %H:%M UTC")

    buyers = buyers or []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONTS["ko"]
    style.font.size = Pt(10)

    section = doc.sections[0]
    apply_page_margins(section, cover=False)
    set_header_footer(section, company, product_name, date_line)

    # ── 표지줄 ──
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(company)
    cr.bold = True
    cr.font.size = Pt(11)
    cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr_ = tp.add_run("바이어 발굴 보고서")
    tr_.bold = True
    tr_.font.size = Pt(28)
    tr_.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr_ = sp.add_run(product_name)
    sr_.font.size = Pt(18)
    sr_.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(date_line)
    dr.font.size = Pt(12)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")

    # ── 섹션 1: 전체 바이어 목록 ──
    _sec_hdr(doc, f"1. 전체 바이어 목록 ({len(buyers)}개)")

    if buyers:
        cols = ["순위", "기업명", "국가", "분류", "이메일 / 웹사이트"]
        list_tbl = doc.add_table(rows=1 + len(buyers), cols=len(cols))
        list_tbl.style = "Table Grid"
        for i, h in enumerate(cols):
            c = list_tbl.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(c, COLORS["navy"])
        for ri, buyer in enumerate(buyers, start=1):
            contact = (
                buyer.get("email") or
                buyer.get("website") or
                buyer.get("contact") or "—"
            )
            row = list_tbl.rows[ri]
            row.cells[0].text = str(ri)
            row.cells[1].text = str(buyer.get("company") or buyer.get("name") or "—")
            row.cells[2].text = str(buyer.get("country") or "—")
            row.cells[3].text = str(buyer.get("type") or buyer.get("category") or buyer.get("classification") or "—")
            row.cells[4].text = str(contact)[:60]
    else:
        doc.add_paragraph("바이어 데이터 없음.")

    doc.add_paragraph("")

    # ── 섹션 2: Top 3 상세 프로파일 ──
    top3 = buyers[:3]
    _sec_hdr(doc, "2. Top 3 바이어 상세 프로파일")

    for rank, buyer in enumerate(top3, start=1):
        company_name = str(buyer.get("company") or buyer.get("name") or f"바이어 {rank}")
        _sub_hdr(doc, f"#{rank}  {company_name}")

        # 기업 개요 박스 (회색 배경)
        overview_tbl = doc.add_table(rows=1, cols=1)
        overview_tbl.style = "Table Grid"
        ov_cell = overview_tbl.rows[0].cells[0]
        set_cell_shading(ov_cell, COLORS["light_gray"])
        ov_cell.text = ""
        ov_lines = []
        if buyer.get("country"):
            ov_lines.append(f"국가: {buyer['country']}")
        if buyer.get("type") or buyer.get("category"):
            ov_lines.append(f"분류: {buyer.get('type') or buyer.get('category')}")
        if buyer.get("revenue") or buyer.get("annual_revenue"):
            ov_lines.append(f"매출: {buyer.get('revenue') or buyer.get('annual_revenue')}")
        if buyer.get("employees") or buyer.get("employee_count"):
            ov_lines.append(f"직원: {buyer.get('employees') or buyer.get('employee_count')}")
        ov_p = ov_cell.paragraphs[0]
        ov_r = ov_p.add_run("\n".join(ov_lines) if ov_lines else "정보 없음")
        ov_r.font.size = Pt(10)

        doc.add_paragraph("")

        # 추천 근거
        reasons = buyer.get("reasons") or buyer.get("rationale") or buyer.get("recommendation_reasons") or []
        if reasons:
            _sub_hdr(doc, "추천 근거")
            if isinstance(reasons, list):
                for rs in reasons:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.add_run(f"• {rs}")
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.add_run(f"• {reasons}")

        # 제품 포트폴리오
        portfolio = buyer.get("portfolio") or buyer.get("products") or buyer.get("product_portfolio") or []
        if portfolio:
            _sub_hdr(doc, "제품 포트폴리오")
            if isinstance(portfolio, list):
                for prod in portfolio[:5]:
                    doc.add_paragraph(f"  · {prod}")
            else:
                doc.add_paragraph(f"  · {portfolio}")

        # 연락처
        _sub_hdr(doc, "연락처")
        contact_info = []
        if buyer.get("address"):
            contact_info.append(f"주소: {buyer['address']}")
        if buyer.get("email"):
            contact_info.append(f"이메일: {buyer['email']}")
        if buyer.get("website"):
            contact_info.append(f"웹사이트: {buyer['website']}")
        if buyer.get("phone"):
            contact_info.append(f"전화: {buyer['phone']}")
        for ci in contact_info:
            doc.add_paragraph(f"  {ci}")
        if not contact_info:
            doc.add_paragraph("  연락처 정보 없음")

        # 참고 출처
        sources = buyer.get("sources") or buyer.get("references") or []
        if sources:
            _sub_hdr(doc, "참고 출처")
            if isinstance(sources, list):
                for s in sources[:3]:
                    doc.add_paragraph(f"  · {s}")
            else:
                doc.add_paragraph(f"  · {sources}")

        doc.add_paragraph("")

    if not top3:
        doc.add_paragraph("상세 프로파일 데이터 없음.")

    # ── 섹션 3: 방법론 및 면책 ──
    _sec_hdr(doc, "3. 방법론 및 면책")

    method_p = doc.add_paragraph(
        "본 바이어 목록은 공개 데이터베이스, 무역 통계, AI 기반 검색을 통해 수집된 잠재 파트너 후보입니다. "
        "실제 거래 전 현지 에이전트를 통한 검증과 최신 기업 정보 확인을 권장합니다."
    )
    method_p.runs[0].font.size = Pt(9)
    method_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"Saudi Pharma Crawler | {timestamp}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    filename = f"sa_03_buyers_{date_str}.docx"
    output_path = out_dir / filename
    doc.save(str(output_path))
    logger.info("P3 report saved: %s", output_path)
    return output_path
